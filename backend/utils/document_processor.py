"""
Document Processor Module
Handles text extraction from various file formats with enhanced OCR support.
"""

import os
import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
from config import Config
import chardet
import io


class DocumentProcessor:

    @staticmethod
    def detect_encoding(file_path):
        """Detect file encoding"""
        with open(file_path, 'rb') as f:
            result = chardet.detect(f.read())
            return result['encoding']

    # -------------------------------------------------------------------------
    # Image preprocessing for better OCR accuracy
    # -------------------------------------------------------------------------

    @staticmethod
    def preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
        """
        Preprocess a PIL image to significantly improve Tesseract OCR accuracy.
        Steps:
          1. Convert to RGB (handles PNGs with alpha channel)
          2. Convert to grayscale
          3. Upscale to at least 1500px wide (improves small-text recognition)
          4. Enhance contrast and apply light sharpening
          5. Convert to pure B&W (binarize) to remove noise
        """
        # 1. Ensure RGB first (RGBA/P mode images need conversion)
        if image.mode in ('RGBA', 'P', 'LA'):
            image = image.convert('RGB')

        # 2. Grayscale
        image = image.convert('L')

        # 3. Upscale if too small — Tesseract needs ~300 DPI; target min 1500px wide
        min_width = 1500
        if image.width < min_width:
            scale = min_width / image.width
            new_size = (int(image.width * scale), int(image.height * scale))
            image = image.resize(new_size, Image.LANCZOS)

        # 4. Enhance contrast
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)

        # 5. Sharpen
        image = image.filter(ImageFilter.SHARPEN)

        # 6. Binarize (threshold to pure black/white)
        image = image.point(lambda x: 0 if x < 140 else 255, '1').convert('L')

        return image

    # -------------------------------------------------------------------------
    # Extractors
    # -------------------------------------------------------------------------

    @staticmethod
    def extract_text_from_pdf(file_path):
        """
        Extract text from PDF using PyMuPDF.
        For each page that has no selectable text (scanned pages), falls back to
        rendering the page as a high-res image and running Tesseract OCR on it.
        """
        try:
            text_parts = []
            doc = fitz.open(file_path)
            total_pages = len(doc)
            ocr_pages = 0

            for page_num in range(total_pages):
                page = doc[page_num]
                page_text = page.get_text().strip()

                # If PyMuPDF returned almost nothing, this is likely a scanned page
                if len(page_text) < 50:
                    try:
                        # Render at 2× scale (~144 DPI) for good OCR quality
                        mat = fitz.Matrix(2.0, 2.0)
                        pix = page.get_pixmap(matrix=mat)
                        img_bytes = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_bytes))
                        img = DocumentProcessor.preprocess_image_for_ocr(img)
                        page_text = pytesseract.image_to_string(
                            img,
                            config='--oem 3 --psm 6'
                        ).strip()
                        if page_text:
                            ocr_pages += 1
                    except Exception as ocr_err:
                        print(f"    ⚠️  OCR failed for page {page_num + 1}: {ocr_err}")
                        page_text = ""

                if page_text:
                    text_parts.append(f"[Page {page_num + 1} of {total_pages}]\n{page_text}")

            doc.close()

            if ocr_pages:
                print(f"    🔍 OCR applied to {ocr_pages} scanned page(s)")

            return "\n\n".join(text_parts).strip()

        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""

    @staticmethod
    def extract_text_from_docx(file_path):
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""

    @staticmethod
    def extract_text_from_txt(file_path):
        """Extract text from TXT file"""
        try:
            encoding = DocumentProcessor.detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding or 'utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return ""

    @staticmethod
    def extract_text_from_image(file_path):
        """
        Extract text from image using OCR with preprocessing.
        Preprocessing steps dramatically improve results for:
          - Scanned ID cards (Aadhaar, PAN, passports)
          - Low-resolution images
          - Images with coloured / noisy backgrounds
        """
        try:
            image = Image.open(file_path)
            processed = DocumentProcessor.preprocess_image_for_ocr(image)

            # Try multi-language first (eng + hin for Indian documents), fallback to eng
            try:
                text = pytesseract.image_to_string(
                    processed,
                    lang='eng+hin',
                    config='--oem 3 --psm 6'
                ).strip()
            except pytesseract.TesseractError:
                text = pytesseract.image_to_string(
                    processed,
                    lang='eng',
                    config='--oem 3 --psm 6'
                ).strip()

            if not text:
                # Last resort: try without preprocessing
                text = pytesseract.image_to_string(
                    image,
                    config='--oem 3 --psm 3'
                ).strip()

            return text if text else f"[Image file: {os.path.basename(file_path)} — no text detected by OCR]"

        except Exception as e:
            print(f"Error extracting text from image: {e}")
            return f"[Image file: {os.path.basename(file_path)} — OCR error: {e}]"

    @staticmethod
    def extract_text_from_markdown(file_path):
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Error extracting text from Markdown: {e}")
            return ""

    @staticmethod
    def extract_text_from_code(file_path):
        """Extract text from code files"""
        try:
            encoding = DocumentProcessor.detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding or 'utf-8', errors='ignore') as f:
                code = f.read()
                # Add language identifier based on extension
                ext = os.path.splitext(file_path)[1][1:]
                return f"```{ext}\n{code}\n```"
        except Exception as e:
            print(f"Error extracting text from code file: {e}")
            return ""

    # -------------------------------------------------------------------------
    # Unified interface
    # -------------------------------------------------------------------------

    @staticmethod
    def process_document(file_path):
        """
        Unified document processing interface.
        Returns: dict with metadata and extracted text.
        """
        filename = os.path.basename(file_path)
        file_ext = os.path.splitext(filename)[1].lower()[1:]

        print(f"Processing document: {filename} (type: {file_ext})")

        text = ""
        doc_type = "unknown"
        page_count = None

        if file_ext == 'pdf':
            text = DocumentProcessor.extract_text_from_pdf(file_path)
            doc_type = "pdf"
            # Estimate page count from [Page N of M] markers
            try:
                import re
                matches = re.findall(r'\[Page \d+ of (\d+)\]', text)
                page_count = int(matches[0]) if matches else None
            except Exception:
                pass

        elif file_ext in ['docx', 'doc']:
            text = DocumentProcessor.extract_text_from_docx(file_path)
            doc_type = "docx"

        elif file_ext == 'txt':
            text = DocumentProcessor.extract_text_from_txt(file_path)
            doc_type = "text"

        elif file_ext in ['md', 'markdown']:
            text = DocumentProcessor.extract_text_from_markdown(file_path)
            doc_type = "markdown"

        elif file_ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']:
            text = DocumentProcessor.extract_text_from_image(file_path)
            doc_type = "image"

        elif file_ext in ['py', 'js', 'java', 'cpp', 'c', 'html', 'css', 'json', 'xml']:
            text = DocumentProcessor.extract_text_from_code(file_path)
            doc_type = "code"

        else:
            # Try as text file
            text = DocumentProcessor.extract_text_from_txt(file_path)
            doc_type = "text"

        result = {
            "filename": filename,
            "file_path": file_path,
            "file_type": doc_type,
            "extension": file_ext,
            "text": text,
            "word_count": len(text.split()),
            "char_count": len(text),
            "page_count": page_count,
            "success": len(text) > 10  # Require at least 10 chars — empty OCR fails gracefully
        }

        print(f"  → Extracted {result['word_count']} words, {result['char_count']} chars"
              + (f", {page_count} pages" if page_count else ""))

        return result


if __name__ == "__main__":
    # Test document processing
    import sys

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        result = DocumentProcessor.process_document(file_path)
        print(f"\nExtracted Text (first 500 chars):\n{result['text'][:500]}...")
        print(f"\nWord count : {result['word_count']}")
        print(f"Page count : {result.get('page_count', 'N/A')}")
        print(f"Success    : {result['success']}")
    else:
        print("Usage: python document_processor.py <file_path>")
